"""Сборка корпуса базы знаний: data/orig_docs -> data/kb/<category>/<file>.

Блок 5.5 требует корпус из 50+ документов в минимум двух (в идеале — четырёх)
форматах: PDF, DOCX, HTML, MD. Исходники лежат в `data/orig_docs` (домен ФТС,
форматы PDF/DOCX), HTML/MD в нужном объёме отсутствуют — поэтому часть DOCX
конвертируется в Markdown, а часть PDF — в минималистичный HTML.

Итог детерминирован (сортировка по пути): повторный запуск даёт тот же корпус.
Категория берётся из пути `orig_docs/ФТ_2023/<category>/...` — это бесплатный
источник метаданных для фильтрации (department/category в ingestion).

Запуск:
    uv run python scripts/prepare_corpus.py
    uv run python scripts/prepare_corpus.py --limit-pdf 32 --limit-docx 20

Скрипт только готовит файлы; индексацию делает scripts/ingest.py.
"""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

# Windows-консоль может быть в cp1251 — переключаем stdout на UTF-8,
# чтобы print с кириллицей не падал/не «кракозябрился».
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = Path(__file__).resolve().parent.parent
ORIG = ROOT / "data" / "orig_docs"
KB = ROOT / "data" / "kb"
INVENTORY = ROOT / "docs" / "data_inventory.md"

# Сколько файлов конвертируем: DOCX -> MD и PDF -> HTML. Итог =
# limit_pdf PDF + limit_docx DOCX + (до N_MD) MD + (до N_HTML) HTML.
N_MD = 8
N_HTML = 8


# Таксономия ПС: русская папка исходников -> латинский slug (см. техдолг
# category-taxonomy). Пересечения («Малахит - Тарифы», «Постконтроль -
# Тарифы») и неизвестное -> raznoe.
CATEGORY_SLUG: dict[str, str] = {
    "Тарифы": "tarify",
    "Малахит": "malahit",
    "Разное": "raznoe",
    "Малахит - Тарифы": "raznoe",
    "Малахит_-_Тарифы": "raznoe",
    "Постконтроль": "postkontrol",
    "Постконтроль - Тарифы": "raznoe",
    "Таможня и Право": "tamozhnya_pravo",
}


def category_from_orig(path: Path) -> str:
    """`orig_docs/ФТ_2023/Малахит/.../file.pdf` -> `malahit`.

    Категория — папка уровня сразу под годовой (`ФТ_2023`); русское имя
    мапится на slug ПС, пересечения и неизвестное — `raznoe`.
    """
    rel = path.relative_to(ORIG)
    parts = rel.parts
    if len(parts) >= 2:
        return CATEGORY_SLUG.get(parts[1], "raznoe")
    return "raznoe"


def _unique_target(category: str, filename: str, seen: dict[str, int]) -> Path:
    """Возвращает свободный путь data/kb/<category>/<filename>, разводит коллизии."""
    target_dir = KB / category
    target_dir.mkdir(parents=True, exist_ok=True)
    stem, suffix = Path(filename).stem, Path(filename).suffix
    candidate = target_dir / filename
    n = seen.get(f"{category}/{filename}", 0)
    while candidate.exists() or n > 0:
        n += 1
        candidate = target_dir / f"{stem}_{n}{suffix}"
    seen[f"{category}/{filename}"] = n
    return candidate


def docx_to_md(src: Path) -> str:
    """DOCX -> Markdown: абзацы + заголовки по стилю + простые таблицы.

    Возвращает пустую строку, если в документе нет извлекаемого текста
    (например, чистая схема/диаграмма) — такой файл в корпус не попадает.
    """
    from docx import Document  # python-docx

    doc = Document(str(src))
    blocks: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            blocks.append(f"# {text}")
        else:
            blocks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            blocks.append(" | ".join(cells))
    return "\n\n".join(blocks).strip()


def pdf_to_html(src: Path) -> str:
    """PDF -> минималистичный HTML: по <p> на страницу (HTMLTagReader его съест).

    Возвращает пустую строку, если ни на одной странице нет текстового слоя
    (скан без OCR) — такой файл в корпус не попадает.
    """
    import pymupdf

    doc = pymupdf.open(str(src))
    pages: list[str] = []
    try:
        for page in doc:
            txt = page.get_text().strip()
            if txt:
                pages.append(txt)
    finally:
        doc.close()
    if not pages:
        return ""

    title = src.stem
    parts = [
        "<html><head><meta charset='utf-8'>",
        f"<title>{title}</title></head><body>",
        f"<h1>{title}</h1>",
    ]
    parts.extend(f"<p>{txt}</p>" for txt in pages)
    parts.append("</body></html>")
    return "\n".join(parts)


def gather() -> tuple[list[Path], list[Path]]:
    """Возвращает (pdf, docx) — отсортированные списки исходных файлов."""
    pdf = sorted({p for p in ORIG.rglob("*") if p.is_file() and p.suffix.lower() == ".pdf"})
    docx = sorted({p for p in ORIG.rglob("*") if p.is_file() and p.suffix.lower() == ".docx"})
    return pdf, docx


def main() -> None:
    parser = argparse.ArgumentParser(description="Сборка корпуса data/kb из data/orig_docs")
    parser.add_argument("--limit-pdf", type=int, default=32)
    parser.add_argument("--limit-docx", type=int, default=20)
    args = parser.parse_args()

    if not ORIG.exists():
        raise SystemExit(f"{ORIG} не найден — сначала положите исходники ФТС.")

    # Чистим предыдущий корпус, чтобы не копить устаревшие файлы.
    if KB.exists():
        shutil.rmtree(KB)
    KB.mkdir(parents=True, exist_ok=True)

    pdf, docx = gather()
    seen: dict[str, int] = {}
    inventory: list[dict] = []

    def register(src: Path, target: Path, fmt: str) -> None:
        inventory.append(
            {
                "file": target.name,
                "category": target.parent.name,
                "format": fmt,
                "size": target.stat().st_size,
            }
        )

    # 1. PDF как есть.
    for p in pdf[: args.limit_pdf]:
        tgt = _unique_target(category_from_orig(p), p.name, seen)
        shutil.copy2(p, tgt)
        register(p, tgt, "pdf")

    # 2. PDF -> HTML (сканируем вперёд, пропуская сканы без текстового слоя).
    html_made = 0
    for p in pdf[args.limit_pdf :]:
        if html_made >= N_HTML:
            break
        content = pdf_to_html(p)
        if not content:
            print(f"  [skip] нет текстового слоя (скан): {p.name}")
            continue
        tgt = _unique_target(category_from_orig(p), p.with_suffix(".html").name, seen)
        tgt.write_text(content, encoding="utf-8")
        register(p, tgt, "html")
        html_made += 1

    # 3. DOCX как есть.
    for d in docx[: args.limit_docx]:
        tgt = _unique_target(category_from_orig(d), d.name, seen)
        shutil.copy2(d, tgt)
        register(d, tgt, "docx")

    # 4. DOCX -> MD (сканируем вперёд, пропуская пустые — схемы/диаграммы).
    md_made = 0
    for d in docx[args.limit_docx :]:
        if md_made >= N_MD:
            break
        content = docx_to_md(d)
        if not content:
            print(f"  [skip] нет извлекаемого текста: {d.name}")
            continue
        tgt = _unique_target(category_from_orig(d), d.with_suffix(".md").name, seen)
        tgt.write_text(content, encoding="utf-8")
        register(d, tgt, "md")
        md_made += 1

    # Итоги + docs/data_inventory.md
    by_fmt: dict[str, int] = {}
    total_size = 0
    for item in inventory:
        by_fmt[item["format"]] = by_fmt.get(item["format"], 0) + 1
        total_size += item["size"]

    lines = [
        "# Инвентаризация корпуса базы знаний",
        "",
        f"Сгенерировано `scripts/prepare_corpus.py` из `data/orig_docs` "
        f"(домен ФТС). Категории — папки верхнего уровня исходников.",
        "",
        "## Сводка",
        "",
        f"- Всего документов: **{len(inventory)}**",
        f"- Общий размер: **{total_size / 1024 / 1024:.1f} МБ**",
        "",
        "| Формат | Число файлов |",
        "|--------|-------------:|",
    ]
    for fmt in ("pdf", "docx", "md", "html"):
        lines.append(f"| {fmt.upper()} | {by_fmt.get(fmt, 0)} |")

    lines += [
        "",
        "## Разбивка по категориям",
        "",
        "| Категория | Число файлов |",
        "|-----------|-------------:|",
    ]
    by_cat: dict[str, int] = {}
    for item in inventory:
        by_cat[item["category"]] = by_cat.get(item["category"], 0) + 1
    for cat, n in sorted(by_cat.items(), key=lambda kv: -kv[1]):
        lines.append(f"| {cat} | {n} |")

    INVENTORY.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(f"Корпус собран в {KB}: {len(inventory)} файлов")
    print("Форматы:", ", ".join(f"{k}={v}" for k, v in sorted(by_fmt.items())))
    print(f"Инвентаризация: {INVENTORY}")


if __name__ == "__main__":
    main()
