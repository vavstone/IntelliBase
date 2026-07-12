"""Конвертация загруженных медиа в content-part для OpenAI Chat Completions.

Native multimodal: изображения отдаются как `image_url`-part (data: base64
URI) прямо в основной chat.completions.create — без отдельного Vision-вызова.
Голос идёт через Whisper-1 (без FFmpeg, Whisper принимает ogg/m4a/mp3
напрямую). PDF/DOCX — извлечение текста через pypdf / python-docx.

Замечание: для image-part `part["image_url"]["url"]` содержит полный base64
изображения. Если этот dict сохраняется в `ChatMessage.media_refs.part`, он
будет загружен заново при каждом обращении к истории. Для учебной задачи —
это допустимо; продакшен потребовал бы внешнего blob-storage с ID.
"""

import base64
from io import BytesIO

from docx import Document
from fastapi import UploadFile
from openai import AsyncOpenAI
from pypdf import PdfReader

ContentPart = dict  # тип-алиас под OpenAI Chat Completions content-part


async def media_to_part(
    media: UploadFile, llm_client: AsyncOpenAI
) -> ContentPart:
    """Конвертирует загруженное медиа в content-part для chat.completions.

    Поддерживаемые типы:
    - `image/*` → image_url-part с data: base64 URI.
    - `audio/*` (а также `application/ogg` — Telegram voice) → Whisper-1
      транскрипт → text-part с пометкой '[пользователь сказал голосом]:'.
    - `application/pdf` → pypdf → text-part '[документ PDF]:'.
    - `application/vnd.openxmlformats-officedocument.wordprocessingml.document`
      → python-docx → text-part '[документ DOCX]:'.

    Raises:
        ValueError: для неподдерживаемых MIME-типов.
    """
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(
            data, media.filename or "audio.ogg", llm_client,
        )
        return {
            "type": "text",
            "text": f"[пользователь сказал голосом]:\n{transcript}",
        }

    if mime == "application/pdf":
        return {
            "type": "text",
            "text": f"[документ PDF]:\n{extract_pdf_text(data)[:30_000]}",
        }

    if mime.endswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{extract_docx_text(data)[:30_000]}",
        }

    raise ValueError(f"Unsupported media type: {mime}")


async def whisper_transcribe(
    audio_bytes: bytes, filename: str, llm_client: AsyncOpenAI,
) -> str:
    """Whisper-1 принимает ogg/m4a/mp3/wav/flac/webm без конвертации."""
    f = BytesIO(audio_bytes)
    f.name = filename  # OpenAI SDK ориентируется на расширение из .name
    result = await llm_client.audio.transcriptions.create(
        model="whisper-1", file=f,
    )
    return result.text


def extract_pdf_text(data: bytes, max_pages: int = 50) -> str:
    """Извлекает текст из PDF. Сканы (мало текста на много страниц) — заглушка."""
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            break
        parts.append(page.extract_text() or "")
    text = "\n\n".join(parts).strip()
    if len(text) < 100 and len(reader.pages) >= 5:
        return "[это скан, OCR пока не поддерживается]"
    return text


def extract_docx_text(data: bytes) -> str:
    """Извлекает текст параграфов и таблиц из DOCX."""
    doc = Document(BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
